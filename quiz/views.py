from django.shortcuts import render
from django.http import HttpResponseBadRequest
from django.views.decorators.csrf import csrf_protect
from .models import Answer, Question, Quiz, StudentAnswer
from course.models import Major, Subject
from rest_framework.decorators import api_view
from django.http import JsonResponse
from user.models import Student, Teacher
from .forms import QuizForm
import docx


# Create your views here.

majors = Major.objects.all()
context = {'majors': majors}

@api_view(['GET'])
def subject_quizzes(request):
    if request.session.has_key('user_id'):
        name = request.session['name']
        context['name'] = name
        context['avatar'] = request.session['avatar']
        
        subjects = Subject.objects.all()
        context['subjects'] = subjects

        return render(request, 'student/subject_quizzes.html', context)
    
    return render(request, 'login/login.html', {'msg':''})


@api_view(['GET'])
def show_quizzes(request, id):
    if request.session.has_key('user_id'):
        name = request.session['name']
        context['name'] = name
        context['avatar'] = request.session['avatar']
        
        subject = Subject.objects.get(pk=id)
        quizzes = subject.quizzes.all()
        context['quizzes'] = quizzes
        context['subject'] = subject

        return render(request, 'student/quizzes.html', context)
    
    return render(request, 'login/login.html', {'msg':''})


@api_view(['GET'])
def quiz_detail(request, id):
    if request.session.has_key('user_id'):
        name = request.session['name']
        context['name'] = name
        context['avatar'] = request.session['avatar']

        quiz = Quiz.objects.get(pk=id)
        context['quiz'] = quiz

        return render(request, 'student/quiz_detail.html', context)
    
    return render(request, 'login/login.html', {'msg':''})


def quiz_data(request, id):
    if request.session.has_key('user_id'):
        name = request.session['name']
        context['name'] = name
        context['avatar'] = request.session['avatar']

        quiz = Quiz.objects.get(pk=id)
        questions = []
        for q in quiz.get_question():
            answers = []
            for a in q.get_answers():
                answers.append(a.text)
            questions.append({str(q): answers})

        return JsonResponse ({
            'data': questions,
            'time': quiz.time,
        })
    
    return render(request, 'login/login.html', {'msg':''})


def save_quiz(request, id):
    if request.is_ajax():
        questions = []
        data = request.POST
        data_ = dict(data.lists())

        data_.pop('csrfmiddlewaretoken')
        quiz = Quiz.objects.get(pk=id)

        for k in data_.keys():
            question = Question.objects.get(quiz=quiz, text=k)
            questions.append(question)

        student = Student.objects.get(pk=request.session['student_id'])

        score = 0
        multiplier = 10 / quiz.number_of_question
        results = []
        correct_answer = None

        for q in questions:
            a_selected = request.POST.get(q.text)
            
            if a_selected != "":
                question_answers = Answer.objects.filter(question=q)
                for a in question_answers:
                    if a_selected == a.text:
                        if a.correct:
                            score += 1
                            correct_answer = a.text
                    else:
                        if a.correct:
                            correct_answer = a.text
                results.append({str(q): {'correct_answer': correct_answer, 'answered': a_selected}})

            else:
                results.append({str(q): 'not answered'})
            
        score_ = score * multiplier
        StudentAnswer.objects.create(quiz=quiz, student=student, score=score_)

        if score_ >= quiz.required_score_to_pass:
            return JsonResponse({
                'passed': True, 
                'score': score_, 
                'results': results
            })
        else:
            return JsonResponse({
                'passed': False,
                'score': score_, 
                'results': results
            })

@api_view(['GET'])
def teacher_manage_score(request):
    if request.session.has_key('user_id'):
        teacher = Teacher.objects.get(pk=request.session['teacher_id'])
        subjects = Subject.objects.filter(major=teacher.major_id)

        contextTeacher['avatar'] = teacher.avatar.url
        contextTeacher['name'] = teacher.name
        contextTeacher['subjects'] = subjects

        return render(request, 'teacher/manage_score.html', contextTeacher)

    return render(request, 'login/login.html', {'msg':''})


def teacher_get_quiz_data(request, id):
    subject = Subject.objects.get(pk=id)
    quizzes = subject.quizzes.all()
    data = []
    for q in quizzes:
        data.append({q.name: q.id})
    print(data)
    return JsonResponse({
        'data': data
    })
    

def get_list_score(request, id):
    if request.session.has_key('user_id'):
        name = request.session['name']
        context['name'] = name
        context['avatar'] = request.session['avatar']

        quiz = Quiz.objects.get(pk=id)
        scores = quiz.quiz_result.all()
        subject = quiz.subject.name

        context['subject'] = subject
        context['quiz'] = quiz
        context['scores'] = scores

        return render(request, 'teacher/view_score.html', context)

    return render(request, 'login/login.html', {'msg':''})
    

contextTeacher = {}
def teacher_get_subject(request):
    if request.session.has_key('user_id'):
        teacher = Teacher.objects.get(pk=request.session['teacher_id'])
        subjects = Subject.objects.filter(major=teacher.major_id)

        contextTeacher['avatar'] = teacher.avatar.url
        contextTeacher['name'] = teacher.name
        contextTeacher['subjects'] = subjects

        return render(request, 'teacher/manage_quiz.html', contextTeacher)

    return render(request, 'login/login.html', {'msg':''})


def quizzes_json(request, id):
    teacher = request.session['name']
    subject = Subject.objects.get(pk=id)
    quizzes = []
    for q in subject.quizzes.all():
        # if q.teacher_id.name == teacher:
        quizzes.append({str(q): q.id})

    print(quizzes)
    return JsonResponse({
        'quizzes': quizzes
    })


def teacher_add_quiz(request, id):
    if request.session.has_key('user_id'):
        subject = Subject.objects.get(pk=id)
        teacher = Teacher.objects.get(name=request.session['name'])

        contextTeacher['avatar'] = teacher.avatar.url
        contextTeacher['name'] = teacher.name
        contextTeacher['subject'] = subject

        # if request.method == 'POST':
        #     name = request.POST['quizName']
        #     numberQuestion = request.POST['numberQuestion']
        #     time = request.POST['time']
        #     requiredScore = request.POST['requiredScore']
        #     difficulity = request.POST['difficulity']
        #     showResult = request.POST['showResult']

        #     Quiz.objects.create(subject=subject, teacher=teacher, name=name, number_of_question=numberQuestion,
        #     time=time, required_score_to_pass=requiredScore, difficulity=difficulity, show_result=showResult)
        #     return render(request, 'teacher/add_question.html', contextTeacher)

        return render(request, 'teacher/add_quiz.html', contextTeacher)
    
    return render(request, 'login/login.html', {'msg':''})


def teacher_save_quiz(request, id):
    if request.is_ajax():
        subject = Subject.objects.get(pk=id)
        teacher = Teacher.objects.get(name=request.session['name'])
        
        data = request.POST
        data_ = dict(data.lists())
        data_.pop('csrfmiddlewaretoken')
        name = data_['name'][0]
        numberQuestion = data_['numberQuestion'][0]
        time = data_['time'][0]
        requiredScore = data_['requiredScore'][0]
        difficulity = data_['difficulity'][0]
        showResult = data_['showResult'][0]
        quiz = Quiz.objects.create(subject=subject, teacher=teacher, name=name, number_of_question=numberQuestion,
            time=time, required_score_to_pass=requiredScore, difficulity=difficulity, show_result=showResult)

        return JsonResponse({
            'data': data_,
            'quiz_id': quiz.id
        })


def teacher_add_question_form(request, id):
    if request.session.has_key('user_id'):
        quiz = Quiz.objects.get(pk=id)
        teacher = Teacher.objects.get(name=request.session['name'])

        contextTeacher['avatar'] = teacher.avatar.url
        contextTeacher['name'] = teacher.name
        contextTeacher['quiz'] = quiz
        contextTeacher['subject'] = quiz.subject.name
        return render(request, 'teacher/add-question-form.html', contextTeacher)

    return render(request, 'login/login.html', {'msg': ''})


def teacher_save_question_form(request, id):
    quiz = Quiz.objects.get(pk=id);
    if request.is_ajax():
        data = request.POST
        data_ = dict(data.lists())
        data_.pop('csrfmiddlewaretoken')
        q = Question.objects.create(quiz=quiz, text=data_['question'][0])
        if(data_['correctAns'][0] == 'A'):
            Answer.objects.create(question=q, text=data_['ansA'][0], correct=True)
        else:
            Answer.objects.create(question=q, text=data_['ansA'][0])

        if(data_['correctAns'][0] == 'B'):
            Answer.objects.create(question=q, text=data_['ansB'][0], correct=True)
        else:
            Answer.objects.create(question=q, text=data_['ansB'][0])

        if(data_['correctAns'][0] == 'C'):
            Answer.objects.create(question=q, text=data_['ansC'][0], correct=True)
        else:
            Answer.objects.create(question=q, text=data_['ansC'][0])

        if(data_['correctAns'][0] == 'D'):
            Answer.objects.create(question=q, text=data_['ansD'][0], correct=True)
        else:
            Answer.objects.create(question=q, text=data_['ansD'][0])

        return JsonResponse({
            'data': data_
        })


def teacher_add_question_bydocx(request, id):
    if request.session.has_key('user_id'):
        quiz = Quiz.objects.get(pk=id)
        teacher = Teacher.objects.get(name=request.session['name'])

        contextTeacher['avatar'] = teacher.avatar.url
        contextTeacher['name'] = teacher.name
        contextTeacher['quiz'] = quiz
        contextTeacher['subject'] = quiz.subject.name

        if request.method == 'POST':
            form = QuizForm(request.POST, request.FILES)
            if form.is_valid():
                quiz.docx = request.FILES['docx']
                quiz.save()
                contextTeacher['form'] = form
                return render(request, 'teacher/add-question-bydocx.html', contextTeacher)

        else:
            form = QuizForm
            contextTeacher['form'] = form
        return render(request, 'teacher/add-question-bydocx.html', contextTeacher)
        
    return render(request, 'login/login.html', {'msg': ''})


def question_bydocx_data(request, id):
    quiz = Quiz.objects.get(pk=id)
    if quiz.docx:
        data = docx.Document(quiz.docx)
        # print(len(data.paragraphs))
        questionAnswer = []
        for i in range(0, len(data.paragraphs), 6):
            q = data.paragraphs[i].text
            a = data.paragraphs[i+1].text
            b = data.paragraphs[i+2].text
            c = data.paragraphs[i+3].text
            d = data.paragraphs[i+4].text
            correctAns = data.paragraphs[i+5].text

            question = Question.objects.create(quiz=quiz, text=q)
            if correctAns == 'A' or correctAns == 'a':
                ans = Answer.objects.create(question=question, text=a, correct=True)
            else:
                ans = Answer.objects.create(question=question, text=a)

            if correctAns == 'B' or correctAns == 'b':
                ans = Answer.objects.create(question=question, text=b, correct=True)
            else:
                ans = Answer.objects.create(question=question, text=b)

            if correctAns == 'C' or correctAns == 'c':
                ans = Answer.objects.create(question=question, text=c, correct=True)
            else:
                ans = Answer.objects.create(question=question, text=c)

            if correctAns == 'D' or correctAns == 'd':
                ans = Answer.objects.create(question=question, text=d, correct=True)
            else:
                ans = Answer.objects.create(question=question, text=d)

            questionAnswer.append({q: {'ansA': a, 'ansB': b, 'ansC': c, 'ansD': d, 'correctAns': correctAns}})

        return JsonResponse({
            'data': questionAnswer
        })
    else:
        return JsonResponse({
            'data': 'nothing'
        })


def teacher_delete_quiz(request, id):
    if request.is_ajax():
        quiz = Quiz.objects.get(pk=id)
        if request.method == 'DELETE':
            quiz.delete()
            return JsonResponse({'status': 'Quiz deleted!'})
        return JsonResponse({'status': 'Invalid request'}, status=400)
    else:
        return HttpResponseBadRequest('Invalid request')